#!/usr/bin/env python
# -*- encoding: utf-8 -*-
import sys
sys.stdout.flush()
from flask import Flask, render_template, request, jsonify, send_from_directory
from flask_cors import CORS
import docx
from simplify_docx import simplify
import mammoth
import json
import os
from werkzeug.utils import secure_filename
import traceback
from docx2pdf import convert
from pdf2image import convert_from_path
from PIL import Image
import base64
from io import BytesIO
import tempfile
import pdfplumber
import re
import pythoncom

app = Flask(__name__)
CORS(app)  # Enable CORS for all routes
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0  # Disable caching for development

# Create uploads directory if it doesn't exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

ALLOWED_EXTENSIONS = {'docx'}

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_page_boundaries_from_docx(doc, html_content):
    """
    Extract page boundaries from DOCX using rendered page breaks.
    Returns a list of lists, where each inner list contains paragraph indices for that page.
    """
    try:
        # Extract paragraphs from HTML to get count
        para_pattern = r'(<p>.*?</p>|<h[1-6]>.*?</h[1-6]>|<table>.*?</table>)'
        html_paragraphs = re.findall(para_pattern, html_content, flags=re.DOTALL)

        print(f"  Total HTML paragraphs: {len(html_paragraphs)}")
        print(f"  Total DOCX paragraphs: {len(doc.paragraphs)}")

        # Build page boundaries by detecting page breaks
        page_boundaries = []
        current_page_paras = []

        for para_idx, paragraph in enumerate(doc.paragraphs):
            # Check if this paragraph contains a page break
            has_page_break = False

            # Check for rendered page breaks
            try:
                if hasattr(paragraph, 'rendered_page_breaks') and paragraph.rendered_page_breaks:
                    has_page_break = True
                    print(f"  Paragraph {para_idx} has rendered page break")
            except:
                pass

            # Check for page breaks in runs
            if not has_page_break:
                try:
                    for run in paragraph.runs:
                        if hasattr(run, 'contains_page_break') and run.contains_page_break:
                            has_page_break = True
                            print(f"  Paragraph {para_idx} has page break in run")
                            break
                except:
                    pass

            # If page break found, save current page and start new one
            if has_page_break and current_page_paras:
                page_boundaries.append(current_page_paras)
                print(f"  Page {len(page_boundaries)}: paragraphs {current_page_paras}")
                current_page_paras = []

            # Add current paragraph to current page
            if para_idx < len(html_paragraphs):
                current_page_paras.append(para_idx)

        # Add final page
        if current_page_paras:
            page_boundaries.append(current_page_paras)
            print(f"  Page {len(page_boundaries)}: paragraphs {current_page_paras}")

        # If we only got 1 page, the document might not have explicit page breaks
        if len(page_boundaries) <= 1:
            print("  WARNING: Only 1 page detected from DOCX page breaks")
            return None, html_paragraphs

        print(f"\n  === DOCX-BASED PAGE DISTRIBUTION ===")
        for page_num, page_para_indices in enumerate(page_boundaries):
            print(f"  Page {page_num + 1}: {len(page_para_indices)} paragraphs {page_para_indices}")

        return page_boundaries, html_paragraphs

    except Exception as e:
        print(f"Error extracting DOCX page boundaries: {str(e)}")
        print(traceback.format_exc())
        return None, None

def extract_page_boundaries_from_pdf(pdf_path, html_content):
    """
    Extract text from each PDF page and map HTML paragraphs to pages.
    Returns a list of lists, where each inner list contains paragraph indices for that page.
    """
    try:
        # Extract paragraphs from HTML
        para_pattern = r'(<p>.*?</p>|<h[1-6]>.*?</h[1-6]>|<table>.*?</table>)'
        html_paragraphs = re.findall(para_pattern, html_content, flags=re.DOTALL)

        # Strip HTML tags to get text content of each paragraph
        def strip_tags(html):
            return re.sub(r'<[^>]+>', '', html).strip()

        para_texts = [strip_tags(p) for p in html_paragraphs]

        # Helper function for flexible text matching
        def text_matches_in_page(para_text, page_text):
            """Try multiple strategies to match paragraph text in page text"""
            if not para_text:
                return False

            # Normalize both texts
            para_norm = ' '.join(para_text.split())
            page_norm = ' '.join(page_text.split())

            # Try different sample sizes
            for sample_size in [50, 100, 150, 200]:
                if len(para_norm) < sample_size:
                    sample = para_norm
                else:
                    sample = para_norm[:sample_size]

                if sample and sample in page_norm:
                    return True

            # Try matching just first 20 words
            para_words = para_norm.split()[:20]
            para_sample = ' '.join(para_words)
            if para_sample in page_norm:
                return True

            return False

        # Extract text from each PDF page
        page_boundaries = []
        matched_paras = set()  # Track which paragraphs we've matched

        with pdfplumber.open(pdf_path) as pdf:
            for page_num, page in enumerate(pdf.pages):
                page_text = page.extract_text() or ""

                # Find which paragraphs belong to this page
                page_para_indices = []

                # Check all paragraphs, not just from current_para_idx
                for para_idx in range(len(para_texts)):
                    # Skip if already matched
                    if para_idx in matched_paras:
                        continue

                    para_text_clean = ' '.join(para_texts[para_idx].split())

                    # Skip empty paragraphs
                    if not para_text_clean:
                        continue

                    # Check if this paragraph's text is in the current page
                    if text_matches_in_page(para_text_clean, page_text):
                        page_para_indices.append(para_idx)
                        matched_paras.add(para_idx)

                page_boundaries.append(page_para_indices)
                print(f"  Page {page_num + 1}: paragraphs {page_para_indices}")

        # Find unmatched paragraphs
        all_para_indices = set(range(len(para_texts)))
        unmatched = all_para_indices - matched_paras
        # Remove empty paragraphs from unmatched
        unmatched = {idx for idx in unmatched if para_texts[idx].strip()}

        if unmatched:
            print(f"  WARNING: {len(unmatched)} unmatched paragraphs: {sorted(unmatched)}")

            # Distribute unmatched paragraphs to nearby pages
            for para_idx in sorted(unmatched):
                # Find the best page for this paragraph based on proximity
                best_page = 0
                min_distance = float('inf')

                for page_num, page_para_indices in enumerate(page_boundaries):
                    if not page_para_indices:
                        continue

                    # Calculate distance to closest paragraph in this page
                    for matched_idx in page_para_indices:
                        distance = abs(matched_idx - para_idx)
                        if distance < min_distance:
                            min_distance = distance
                            best_page = page_num

                # Add to best page
                page_boundaries[best_page].append(para_idx)
                page_boundaries[best_page].sort()  # Keep sorted
                print(f"  Assigned unmatched paragraph {para_idx} to page {best_page + 1}")

        # Log final distribution
        print("\n  === FINAL PAGE DISTRIBUTION ===")
        for page_num, page_para_indices in enumerate(page_boundaries):
            print(f"  Page {page_num + 1}: {len(page_para_indices)} paragraphs {page_para_indices}")

        return page_boundaries, html_paragraphs

    except Exception as e:
        print(f"Error extracting PDF page boundaries: {str(e)}")
        print(traceback.format_exc())
        return None, None

def split_html_into_pages(html_content, num_pages, doc=None, pdf_path=None):
    """
    Split HTML content based on DOCX page boundaries if available,
    fall back to PDF, then equal division.
    """
    if num_pages <= 1:
        return [html_content]

    # Try DOCX-based splitting first (most accurate)
    if doc:
        print("Extracting page boundaries from DOCX...")
        page_boundaries, html_paragraphs = extract_page_boundaries_from_docx(doc, html_content)

        if page_boundaries and html_paragraphs:
            # Build pages using the boundary mapping
            pages = []
            for page_indices in page_boundaries:
                if page_indices:
                    page_content = ''.join([html_paragraphs[i] for i in page_indices])
                    pages.append(page_content)
                else:
                    pages.append('')  # Empty page

            print(f"Successfully split HTML using DOCX boundaries: {len(pages)} pages")
            return pages
        else:
            print("DOCX boundary extraction failed or returned 1 page, trying PDF...")

    # Try PDF-based splitting as fallback
    if pdf_path and os.path.exists(pdf_path):
        print("Extracting page boundaries from PDF...")
        page_boundaries, html_paragraphs = extract_page_boundaries_from_pdf(pdf_path, html_content)

        if page_boundaries and html_paragraphs:
            # Build pages using the boundary mapping
            pages = []
            for page_indices in page_boundaries:
                if page_indices:
                    page_content = ''.join([html_paragraphs[i] for i in page_indices])
                    pages.append(page_content)
                else:
                    pages.append('')  # Empty page

            print(f"Successfully split HTML using PDF boundaries: {len(pages)} pages")
            return pages
        else:
            print("PDF boundary extraction failed, falling back to equal division")

    # Fallback: Split HTML by equal division
    print("Using equal division for HTML splitting")
    paragraphs = re.findall(r'(<p>.*?</p>|<h[1-6]>.*?</h[1-6]>|<table>.*?</table>)', html_content, flags=re.DOTALL)

    # Calculate paragraphs per page
    paras_per_page = max(1, len(paragraphs) // num_pages)

    pages = []
    for i in range(num_pages):
        start_idx = i * paras_per_page
        end_idx = start_idx + paras_per_page if i < num_pages - 1 else len(paragraphs)
        page_content = ''.join(paragraphs[start_idx:end_idx])
        pages.append(page_content)

    return pages

def split_json_into_pages(json_str, num_pages, doc=None, pdf_path=None, html_content=None):
    """
    Split JSON content based on DOCX page boundaries if available,
    fall back to PDF, then equal division.
    """
    if num_pages <= 1:
        return [json_str]

    try:
        json_obj = json.loads(json_str)

        # Navigate to the body VALUE array
        if isinstance(json_obj, dict) and 'VALUE' in json_obj:
            body_list = json_obj['VALUE']
            if isinstance(body_list, list) and len(body_list) > 0:
                body = body_list[0]
                if isinstance(body, dict) and 'VALUE' in body:
                    elements = body['VALUE']

                    # Try DOCX-based splitting first
                    if doc and html_content:
                        print("Extracting JSON page boundaries from DOCX...")
                        page_boundaries, html_paragraphs = extract_page_boundaries_from_docx(doc, html_content)

                        if page_boundaries:
                            # Build pages using the boundary mapping
                            pages = []
                            for page_indices in page_boundaries:
                                if page_indices:
                                    # Get JSON elements for this page
                                    page_elements = [elements[i] for i in page_indices if i < len(elements)]
                                    page_obj = {
                                        'TYPE': 'document',
                                        'VALUE': [{
                                            'TYPE': 'body',
                                            'VALUE': page_elements
                                        }]
                                    }
                                    pages.append(json.dumps(page_obj, indent=2))
                                else:
                                    # Empty page
                                    page_obj = {
                                        'TYPE': 'document',
                                        'VALUE': [{
                                            'TYPE': 'body',
                                            'VALUE': []
                                        }]
                                    }
                                    pages.append(json.dumps(page_obj, indent=2))

                            print(f"Successfully split JSON using DOCX boundaries: {len(pages)} pages")
                            return pages
                        else:
                            print("DOCX boundary extraction failed for JSON, trying PDF...")

                    # Try PDF-based splitting as fallback
                    if pdf_path and html_content and os.path.exists(pdf_path):
                        print("Extracting JSON page boundaries from PDF...")
                        page_boundaries, html_paragraphs = extract_page_boundaries_from_pdf(pdf_path, html_content)

                        if page_boundaries:
                            # Build pages using the boundary mapping
                            # Assume JSON elements correspond 1:1 with HTML paragraphs
                            pages = []
                            for page_indices in page_boundaries:
                                if page_indices:
                                    # Get JSON elements for this page
                                    page_elements = [elements[i] for i in page_indices if i < len(elements)]
                                    page_obj = {
                                        'TYPE': 'document',
                                        'VALUE': [{
                                            'TYPE': 'body',
                                            'VALUE': page_elements
                                        }]
                                    }
                                    pages.append(json.dumps(page_obj, indent=2))
                                else:
                                    # Empty page
                                    page_obj = {
                                        'TYPE': 'document',
                                        'VALUE': [{
                                            'TYPE': 'body',
                                            'VALUE': []
                                        }]
                                    }
                                    pages.append(json.dumps(page_obj, indent=2))

                            print(f"Successfully split JSON using PDF boundaries: {len(pages)} pages")
                            return pages
                        else:
                            print("PDF boundary extraction failed for JSON, falling back to equal division")

                    # Fallback: Split elements across pages equally
                    print("Using equal division for JSON splitting")
                    elements_per_page = max(1, len(elements) // num_pages)
                    pages = []

                    for i in range(num_pages):
                        start_idx = i * elements_per_page
                        end_idx = start_idx + elements_per_page if i < num_pages - 1 else len(elements)

                        # Create page structure
                        page_obj = {
                            'TYPE': 'document',
                            'VALUE': [{
                                'TYPE': 'body',
                                'VALUE': elements[start_idx:end_idx]
                            }]
                        }
                        pages.append(json.dumps(page_obj, indent=2))

                    return pages

        # Fallback: return full JSON for all pages
        return [json_str] * num_pages

    except Exception as e:
        print(f"Error splitting JSON: {str(e)}")
        return [json_str] * num_pages

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    try:
        print("Upload request received")  # Debug log

        if 'file' not in request.files:
            print("No file in request")
            return jsonify({'error': 'No file part'}), 400

        file = request.files['file']
        print(f"File received: {file.filename}")

        if file.filename == '':
            return jsonify({'error': 'No selected file'}), 400

        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            filepath = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            print(f"Saving to: {filepath}")
            file.save(filepath)

            try:
                print("Converting to HTML...")
                # Convert to HTML using mammoth
                with open(filepath, "rb") as docx_file:
                    result = mammoth.convert_to_html(docx_file)
                    html_content = result.value
                print(f"HTML conversion complete: {len(html_content)} chars")

                # Debug: Save first 1000 chars of Mammoth HTML
                print(f"\n=== MAMMOTH HTML SAMPLE ===")
                print(html_content[:1000])
                print("=" * 50)

                print("Simplifying DOCX...")
                # Simplify using simplify-docx
                doc = docx.Document(filepath)
                simplified_json = simplify(doc)
                print("Simplification complete")

                # Debug: Show JSON structure sample
                print(f"\n=== JSON STRUCTURE SAMPLE ===")
                print(json.dumps(simplified_json, indent=2)[:1500])
                print("=" * 50)

                print("Detecting page count from DOCX...")
                # Count pages using page breaks in the document
                num_pages = 1
                for para in doc.paragraphs:
                    # Count page breaks in paragraph runs
                    for run in para.runs:
                        if '\f' in run.text or '\x0c' in run.text:
                            num_pages += 1

                # Also check for hard page breaks in the XML
                for section in doc.sections:
                    # Sections often indicate page breaks
                    pass

                print(f"Detected {num_pages} pages from document structure")

                print("Converting DOCX to image...")
                # Convert DOCX to PDF, then to image
                image_data = []
                pdf_path = None
                pdf_created_successfully = False

                try:
                    # Set Poppler path
                    poppler_path = os.path.join(os.getcwd(), "poppler-24.08.0", "Library", "bin")
                    print(f"Using Poppler from: {poppler_path}")

                    # Create a temporary PDF file
                    pdf_path = filepath.replace('.docx', '.pdf')

                    # Initialize COM for Word automation on Windows
                    pythoncom.CoInitialize()
                    try:
                        # Use keep_active=True to prevent closing other Word documents
                        convert(filepath, pdf_path, keep_active=True)
                        print(f"PDF created: {pdf_path}")
                        pdf_created_successfully = True
                    finally:
                        pythoncom.CoUninitialize()

                    # Convert PDF to images with Poppler path - GET ALL PAGES
                    images = convert_from_path(pdf_path, dpi=150, poppler_path=poppler_path)
                    print(f"Converted to {len(images)} page(s)")

                    # Update page count from actual PDF if available
                    if images:
                        num_pages = len(images)
                        print(f"Updated page count from PDF: {num_pages}")

                    # Convert ALL pages to base64 (not just first page)
                    if images:
                        for i, img in enumerate(images):
                            buffered = BytesIO()
                            img.save(buffered, format="PNG")
                            img_base64 = base64.b64encode(buffered.getvalue()).decode('utf-8')
                            image_data.append(img_base64)
                        print(f"Converted {len(image_data)} pages to base64")

                except Exception as img_error:
                    print(f"Warning: Could not convert to image: {str(img_error)}")
                    print(traceback.format_exc())
                    # Keep num_pages from document detection
                    print(f"Using detected page count: {num_pages}")

                # Split HTML and JSON into pages BEFORE cleaning up PDF
                # Pass doc object (primary), pdf_path (fallback), and html_content
                json_str = json.dumps(simplified_json, indent=2)

                if pdf_created_successfully and pdf_path and os.path.exists(pdf_path):
                    html_pages = split_html_into_pages(html_content, num_pages, doc=doc, pdf_path=pdf_path)
                    json_pages = split_json_into_pages(json_str, num_pages, doc=doc, pdf_path=pdf_path, html_content=html_content)
                else:
                    html_pages = split_html_into_pages(html_content, num_pages, doc=doc)
                    json_pages = split_json_into_pages(json_str, num_pages, doc=doc, html_content=html_content)

                # Now clean up PDF and uploaded file
                if pdf_path and os.path.exists(pdf_path):
                    try:
                        os.remove(pdf_path)
                        print(f"Cleaned up PDF: {pdf_path}")
                    except:
                        pass

                try:
                    os.remove(filepath)
                    print(f"Cleaned up uploaded file: {filepath}")
                except:
                    pass

                print(f"Split into {num_pages} pages: {len(html_pages)} HTML, {len(json_pages)} JSON")

                response_data = {
                    'html': html_content,  # Keep full HTML for backward compatibility
                    'json': json_str,  # Keep full JSON for backward compatibility
                    'htmlPages': html_pages,  # NEW: HTML split by pages
                    'jsonPages': json_pages,  # NEW: JSON split by pages
                    'images': image_data,  # Array of page images
                    'pageCount': num_pages,
                    'filename': filename
                }
                print("Sending response")
                return jsonify(response_data)

            except Exception as e:
                print(f"Error processing file: {str(e)}")
                print(traceback.format_exc())
                if os.path.exists(filepath):
                    os.remove(filepath)
                return jsonify({'error': f'Processing error: {str(e)}'}), 500

        return jsonify({'error': 'Invalid file type. Please upload a .docx file'}), 400

    except Exception as e:
        print(f"Unexpected error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Server error: {str(e)}'}), 500

@app.route('/json-to-html', methods=['POST'])
def json_to_html():
    try:
        data = request.get_json()
        json_structure = data.get('json')

        if not json_structure:
            return jsonify({'error': 'No JSON data provided'}), 400

        # Parse JSON string if it's a string
        if isinstance(json_structure, str):
            json_structure = json.loads(json_structure)

        print(f"\n=== JSON to HTML Conversion ===")
        print(f"Input JSON structure (first 200 chars): {str(json_structure)[:200]}")

        # Convert the simplified JSON structure back to HTML
        html = convert_json_to_html(json_structure)

        print(f"Output HTML (first 500 chars): {html[:500]}")
        print(f"Output HTML length: {len(html)}")
        print("=" * 40)

        return jsonify({'html': html})

    except json.JSONDecodeError as e:
        return jsonify({'error': f'Invalid JSON: {str(e)}'}), 400
    except Exception as e:
        print(f"Error converting JSON to HTML: {str(e)}")
        print(traceback.format_exc())
        return jsonify({'error': f'Conversion error: {str(e)}'}), 500

def is_heading(text):
    """Detect if text is likely a heading based on patterns"""
    import re
    if not isinstance(text, str):
        return None

    text = text.strip()

    # H1: Contains keywords like "Framework", "Policy" and ends with colon OR is short and title-cased
    if len(text) < 120:
        if text.endswith(':') and any(word in text for word in ['Framework', 'Plan', 'Strategy', 'Policy', 'Assessment', 'Delivery']):
            return 'h1'
        # Also detect title without colon if it contains these keywords and no period at end
        if not text.endswith('.') and any(word in text for word in ['Framework:', 'Plan:', 'Strategy:', 'Policy:', 'Delivery']):
            return 'h1'

    # H2: Starts with number pattern like "1.0", "2.1", "A.", etc.
    if re.match(r'^(\d+\.)+\d*\s+[A-Z]', text) and len(text) < 100:
        return 'h2'

    # H3: Starts with letter pattern like "a)", "(i)", etc.
    if re.match(r'^[a-z]\)|^\([ivx]+\)', text) and len(text) < 80:
        return 'h3'

    return None

def convert_json_to_html(obj):
    """
    Convert simplified JSON structure back to HTML
    """
    if obj is None:
        return ''

    # Handle string/text directly
    if isinstance(obj, str):
        return obj

    # Handle list of elements
    if isinstance(obj, list):
        return ''.join([convert_json_to_html(item) for item in obj])

    # Handle dictionary/object
    if isinstance(obj, dict):
        # Handle both uppercase (from simplify-docx) and lowercase keys
        tag_type = obj.get('TYPE') or obj.get('type', 'div')
        value = obj.get('VALUE') or obj.get('value', '')
        children = obj.get('children', [])

        # Special handling for text type - just return the text without wrapping
        if tag_type == 'text':
            if isinstance(value, str):
                return value
            elif isinstance(value, (list, dict)):
                return convert_json_to_html(value)
            elif 'text' in obj:
                return obj['text']
            return ''

        # For paragraphs, check if the text content is actually a heading
        if tag_type == 'paragraph':
            # Extract text content to check for heading patterns
            if isinstance(value, list) and len(value) > 0:
                # Get the first text element
                first_elem = value[0]
                if isinstance(first_elem, dict) and first_elem.get('TYPE') == 'text':
                    text_content = first_elem.get('VALUE', '')
                    heading_tag = is_heading(text_content)
                    if heading_tag:
                        # This is a heading, convert the entire value
                        value_html = convert_json_to_html(value)
                        return f'<{heading_tag}>{value_html}</{heading_tag}>'

        # Map simplified types to HTML tags
        tag_map = {
            'document': 'div',
            'body': 'div',
            'paragraph': 'p',
            'table': 'table',
            'table-row': 'tr',
            'table-cell': 'td',
            'heading': 'h2',
            'list': 'ul',
            'list-item': 'li',
        }

        html_tag = tag_map.get(tag_type, 'div')

        # Handle VALUE that contains nested structure (list or dict)
        if isinstance(value, (list, dict)):
            value_html = convert_json_to_html(value)
            return f'<{html_tag}>{value_html}</{html_tag}>'

        # Handle text content
        if value and not children:
            return f'<{html_tag}>{value}</{html_tag}>'

        # Handle children
        if children:
            children_html = convert_json_to_html(children)
            return f'<{html_tag}>{children_html}</{html_tag}>'

        # Handle objects with direct text
        if 'text' in obj:
            return obj['text']

        return ''

    return str(obj)

if __name__ == '__main__':
    app.run(debug=True, host='127.0.0.1', port=5000, use_reloader=False, threaded=True)
